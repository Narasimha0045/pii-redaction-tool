"""Redact common PII from a DOCX using regexes, spaCy, and Faker."""
from __future__ import annotations

import argparse
import ipaddress
import json
import re
from collections import Counter
from pathlib import Path

import spacy
from docx import Document
from faker import Faker

EMAIL = re.compile(r"(?<![\w.+-])[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?![\w-])")
PHONE = re.compile(r"(?<![\d-])(?:\+\d{1,3}[ .-]?)?(?:\d{3}[ .-]\d{3}[ .-]\d{4}|\d{5}[ .-]\d{5})(?![\d-])")
SSN = re.compile(r"(?<!\d)\d{3}-\d{2}-\d{4}(?!\d)")
CARD = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
IPV4 = re.compile(r"(?<![\w.])(?:\d{1,3}\.){3}\d{1,3}(?![\w.])")
IPV6 = re.compile(r"(?<![\w:])(?:[0-9A-Fa-f]{0,4}:){2,7}[0-9A-Fa-f]{0,4}(?![\w:])")
DOB = re.compile(r"\b(?:dob|date of birth|born(?: on)?)\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]+\s+\d{4})", re.I)
ADDRESS = re.compile(r"\b\d{1,5}[A-Za-z-]*\s+[A-Za-z0-9.' -]{2,60}\s+(?:Street|St\.?|Road|Rd\.?|Avenue|Ave\.?|Lane|Ln\.?|Drive|Dr\.?|Boulevard|Blvd\.?|Nagar|Marg)\b(?:,?\s*[A-Za-z][A-Za-z .'-]{1,40}\s+(?:\d{6}|\d{5}(?:-\d{4})?))?", re.I)


def is_valid_card(value: str) -> bool:
    digits = re.sub(r"\D", "", value)
    if not 13 <= len(digits) <= 19 or len(set(digits)) == 1:
        return False

    total = 0
    for index, digit in enumerate(reversed(digits)):
        number = int(digit)
        if index % 2 == 1:
            number = number * 2
            if number > 9:
                number -= 9
        total += number
    return total % 10 == 0


def is_valid_ip(value: str) -> bool:
    try:
        ipaddress.ip_address(value)
        return True
    except ValueError:
        return False


def add_matches(detections: list[dict], entity_type: str, pattern: re.Pattern, text: str, group: int = 0, validator=None) -> None:
    for match in pattern.finditer(text):
        value = match.group(group)
        if validator is None or validator(value):
            start, end = match.span(group)
            detections.append({"type": entity_type, "text": value, "start": start, "end": end})


def remove_overlaps(detections: list[dict]) -> list[dict]:
    """Keep the first detector result when two detectors identify the same text."""
    kept = []
    for detection in detections:
        overlaps = any(detection["start"] < current["end"] and detection["end"] > current["start"] for current in kept)
        if not overlaps:
            kept.append(detection)
    return kept


def find_pii(text: str, nlp) -> list[dict]:
    """Return non-overlapping PII detections for one paragraph of text."""
    detections = []
    add_matches(detections, "EMAIL", EMAIL, text)
    add_matches(detections, "PHONE", PHONE, text)
    add_matches(detections, "SSN", SSN, text)
    add_matches(detections, "CREDIT_CARD", CARD, text, validator=is_valid_card)
    add_matches(detections, "IP_ADDRESS", IPV4, text, validator=is_valid_ip)
    add_matches(detections, "IP_ADDRESS", IPV6, text, validator=is_valid_ip)
    add_matches(detections, "DOB", DOB, text, group=1)
    add_matches(detections, "ADDRESS", ADDRESS, text)

    for entity in nlp(text).ents:
    # Ignore single-word
        if entity.label_ == "PERSON" and len(entity.text.split()) < 2:
            continue
        entity_type = {
            "PERSON": "PERSON",
            "ORG": "COMPANY",
        }.get(entity.label_)
        if (
            entity_type
            and entity.text.upper() not in {"SSN", "DOB", "IP", "CARD"}
            and not any(char.isdigit() for char in entity.text)
        ):
            detections.append(
                {
                    "type": entity_type,
                    "text": entity.text,
                    "start": entity.start_char,
                    "end": entity.end_char,
                }
            )
    return remove_overlaps(detections)

def replacement_for(entity_type: str, original: str, mapping: dict, fake: Faker) -> str:
    key = (entity_type, original.casefold())
    if key in mapping:
        return mapping[key]

    generators = {
        "PERSON": fake.name,
        "EMAIL": fake.safe_email,
        "PHONE": fake.phone_number,
        "COMPANY": fake.company,
        "ADDRESS": lambda: fake.address().replace("\n", ", "),
        "SSN": fake.ssn,
        "CREDIT_CARD": fake.credit_card_number,
        "DOB": lambda: fake.date_of_birth(minimum_age=25, maximum_age=65).strftime("%d/%m/%Y"),
        "IP_ADDRESS": fake.ipv4_public,
    }
    replacement = generators[entity_type]()
    mapping[key] = replacement
    return replacement


def redact_text(text: str, detections: list[dict], mapping: dict, fake: Faker, counts: Counter) -> str:
    for detection in sorted(detections, key=lambda item: item["start"], reverse=True):
        replacement = replacement_for(detection["type"], detection["text"], mapping, fake)
        text = text[:detection["start"]] + replacement + text[detection["end"]:]
        counts[detection["type"]] += 1
    return text


def document_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def write_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def redact_document(input_path: Path, output_path: Path) -> dict:
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError as error:
        raise RuntimeError("Install the spaCy model first: python -m spacy download en_core_web_sm") from error

    fake = Faker("en_US")
    fake.seed_instance(20260724)
    mapping = {}
    counts = Counter()
    document = Document(input_path)

    for paragraph in document_paragraphs(document):
        detections = find_pii(paragraph.text, nlp)
        if detections:
            write_paragraph(paragraph, redact_text(paragraph.text, detections, mapping, fake, counts))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {"output": str(output_path), "counts": dict(counts), "unique_replacements": len(mapping)}


def main() -> None:
    parser = argparse.ArgumentParser(description="Redact PII from a DOCX file.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    print(json.dumps(redact_document(args.input, args.output), indent=2))


if __name__ == "__main__":
    main()
